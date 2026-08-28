import os
import sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Sets inner padding for table cell (in twips, 20 twips = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_callout(doc, text_list, title="KEY ARCHITECTURAL INSIGHT", border_color="4F46E5", bg_color="F8FAFC"):
    """Adds a stylish callout box with a colored left accent border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=160, bottom=160, left=240, right=200)
    
    # Left border only
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.bold = True
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(79, 70, 229)
    
    for t in text_list:
        p_text = cell.add_paragraph()
        p_text.paragraph_format.space_before = Pt(2)
        p_text.paragraph_format.space_after = Pt(2)
        r = p_text.add_run(t)
        r.font.name = "Segoe UI"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def style_table(table, header_bg="1E293B", col_widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Format Header
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_background(cell, header_bg)
        set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.name = "Segoe UI"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)
                
    # Format Body
    for r_idx, row in enumerate(table.rows[1:], start=1):
        bg = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        for c_idx, cell in enumerate(row.cells):
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Segoe UI"
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(51, 65, 85)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

def create_document():
    doc = docx.Document()
    
    # Page setup - 0.75 inch margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)

    # -------------------------------------------------------------
    # COVER / HEADER TITLE
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    t_run = title_p.add_run("SkillMatch AI")
    t_run.font.name = "Segoe UI"
    t_run.font.size = Pt(26)
    t_run.bold = True
    t_run.font.color.rgb = RGBColor(79, 70, 229) # Indigo

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(14)
    s_run = sub_p.add_run("Comprehensive Work & Data Flow Architecture Specification\nStudent & Employer Portals, State Transitions, and UI/UX Revamp Guide")
    s_run.font.name = "Segoe UI"
    s_run.font.size = Pt(13)
    s_run.font.color.rgb = RGBColor(100, 116, 139) # Slate

    # Metadata banner
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(18)
    m_run = meta_p.add_run("Author: Lead AI & Mobile Systems Engineer  |  Target: Mobile UI Engineering Team  |  Status: Complete Architecture Spec")
    m_run.font.size = Pt(8.5)
    m_run.font.italic = True
    m_run.font.color.rgb = RGBColor(148, 163, 184)

    # Divider line
    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(14)
    d_run = divider.add_run("―" * 58)
    d_run.font.color.rgb = RGBColor(226, 232, 240)

    # -------------------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY & SYSTEM TOPOLOGY
    # -------------------------------------------------------------
    h1 = doc.add_heading("1. Executive Summary & Platform Topology", level=1)
    h1.runs[0].font.color.rgb = RGBColor(15, 23, 42)
    
    p = doc.add_paragraph(
        "SkillMatch AI is an intelligent internship and early-career job placement ecosystem designed to eliminate recruitment friction for students and employers. "
        "The platform bridges the gap between unstructured candidate profiles (resumes) and job specifications using a hybrid matching formula combining "
        "ontology-driven skills taxonomy normalization (0.6 weight) and 384-dimensional dense semantic vector embeddings (0.4 weight)."
    )
    
    add_callout(
        doc,
        [
            "• Student Experience: Frictionless resume upload (PDF/DOCX), automatic skill extraction, human-in-the-loop verification, hybrid-scored opportunity discovery feed, and live recruitment status tracking.",
            "• Employer Experience: Rapid job publishing, taxonomy skill auto-tagging, real-time AI-ranked applicant pool review, candidate explainability breakdowns, and 1-click status progression with automated student notifications.",
            "• Data Privacy & Compliance: Full compliance with Republic Act No. 10173 (Philippine Data Privacy Act of 2012) supporting machine-readable data export (/auth/me/export) and permanent account/file erasure (/auth/me)."
        ],
        title="CORE ECOSYSTEM VALUE PILLARS"
    )

    # -------------------------------------------------------------
    # SECTION 2: STUDENT PORTAL WORK & DATA FLOW
    # -------------------------------------------------------------
    h1 = doc.add_heading("2. Student Portal: End-to-End Work & Data Flow", level=1)
    h1.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "The Student Portal guides candidate onboarding from resume ingestion to job matching and interview tracking. Below is the sequential breakdown across screens, UI interactions, API payloads, and data transformations."
    )

    # Step 2.1
    h2 = doc.add_heading("Phase 2.1: Authentication & Session Initialization", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)
    
    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("LoginScreen.tsx / RegisterScreen.tsx\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Student inputs Full Name, Email, Password, and selects role = 'student'. On submission, app authenticates and persists JWT token securely.\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("POST /api/v1/auth/register or POST /api/v1/auth/login\n")
    p.add_run("• Outgoing Payload:\n")
    p.add_run('  { "email": "student@univ.edu", "password": "...", "full_name": "Juan Dela Cruz", "role": "student" }\n').italic = True
    p.add_run("• Response & State Mutation: ").bold = True
    p.add_run("Returns JWT access_token + User object. AuthContext stores token in Expo SecureStore and sets Axios/Fetch default authorization header Bearer <token>.\n")

    # Step 2.2
    h2 = doc.add_heading("Phase 2.2: Resume Ingestion, Extraction & Embedding Generation", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("ResumeUploadScreen.tsx\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Student taps 'Select Document' -> DocumentPicker opens native file selector for PDF or DOCX -> File size (<5MB) verified -> Student taps 'Upload & Parse Resume' -> Animated scanning state renders while backend processes.\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("POST /api/v1/resumes/upload (Content-Type: multipart/form-data)\n")
    p.add_run("• Backend Data Transformation Pipeline:\n")
    p.add_run("  1. Magic Byte Validation: Validates true MIME signatures (puremagic) to prevent renamed malicious payloads.\n")
    p.add_run("  2. Isolated Storage: File saved to storage/resumes/{user_id}/{uuid}.pdf with sanitized paths.\n")
    p.add_run("  3. Text Extraction: PDFPlumber / python-docx extracts raw unstructured textual content.\n")
    p.add_run("  4. LLM Field Extractor: OpenAI/Claude (with heuristic fallback) extracts structured JSON containing skills, education, experience, certifications, summary.\n")
    p.add_run("  5. Taxonomy Normalization: Skills mapped to canonical dictionary (e.g. 'React.js' -> 'React', 'py' -> 'Python').\n")
    p.add_run("  6. 384-Dim Vector Embedding: Sentence-Transformers generates dense semantic vector embedding from synthesized resume text representation.\n")
    p.add_run("• Response:\n")
    p.add_run('  { "id": "res_123", "status": "parsed", "parsed_data": { "skills": ["Python", "FastAPI", "React", "PostgreSQL"], "education": [...], "experience": [...] } }\n').italic = True

    # Step 2.3
    h2 = doc.add_heading("Phase 2.3: Human-in-the-Loop Skill Verification & Editing", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("ResumeReviewScreen.tsx\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Student reviews extracted skills rendered as dismissible chips. Student can remove inaccurate tags, add missing skills via text input, and tap 'Save & Activate Profile'.\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("PUT /api/v1/resumes/{resume_id}/parsed-data\n")
    p.add_run("• Outgoing Payload: ").bold = True
    p.add_run('{ "parsed_data": { "skills": ["Python", "Docker", "FastAPI", "PostgreSQL"], ... }, "status": "active" }\n').italic = True
    p.add_run("• Backend Action: ").bold = True
    p.add_run("Normalizes updated skills list, regenerates the 384-dimensional vector embedding, updates ResumeSkill relational records, and marks status as 'active'.\n")

    # Step 2.4
    h2 = doc.add_heading("Phase 2.4: Recommendation Discovery & Opportunity Matching Feed", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("RecommendationFeedScreen.tsx (Main Tab: 'Opportunities')\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Student browses AI-matched job opportunities. Feed merges Internal Employer Postings with External Synced Jobs (JSearch/RapidAPI). Each card displays match percentage badge, company name, job type chip (Internship / OJT / Full-time), location, and explainability breakdown (matched skills in green, missing skills in amber).\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("GET /api/v1/matches/recommendations\n")
    p.add_run("• Hybrid Scoring Formula (Backend):\n")
    p.add_run("  Final Match Score = (0.60 * Taxonomy Skill Overlap Score) + (0.40 * Cosine Similarity(Resume Vector, Job Vector))\n").italic = True
    p.add_run("• Response Item Schema:\n")
    p.add_run('  {\n'
               '    "match_score": 0.88,\n'
               '    "skill_score": 0.85,\n'
               '    "target_type": "internal",\n'
               '    "target": { "id": "job_01", "title": "Junior Python / Backend Developer", "employer": { "company_name": "Tech Corp" } },\n'
               '    "explanation": {\n'
               '      "matched_skills": ["Python", "FastAPI", "PostgreSQL"],\n'
               '      "missing_skills": ["Docker", "Redis"],\n'
               '      "summary": "You possess 3 of 5 required skills (88% overall alignment)."\n'
               '    }\n'
               '  }\n').italic = True

    # Step 2.5
    h2 = doc.add_heading("Phase 2.5: Application Submission & Live Status Stepper", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("RecommendationFeedScreen.tsx (Action) -> ApplicationTrackerScreen.tsx (View)\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Student taps 'Apply Now' on an internal job card -> Optional cover note modal opens -> On confirm, application record is created -> Student navigates to 'Applications' tab to see visual 5-step progress stepper.\n")
    p.add_run("• API Endpoints:\n")
    p.add_run('  - POST /api/v1/applications/ (Payload: { "job_posting_id": "job_01", "notes": "Eager to contribute." })\n'
               '  - GET /api/v1/applications/my-applications\n')
    p.add_run("• 5-Step Visual Stepper States: ").bold = True
    p.add_run("1. Submitted  →  2. Under Review  →  3. Shortlisted  →  4. Interview Scheduled  →  5. Accepted (or Rejected/Withdrawn).\n")

    # Step 2.6
    h2 = doc.add_heading("Phase 2.6: In-App Status Notification Alerts", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Component / Endpoint: ").bold = True
    p.add_run("GET /api/v1/notifications/ and GET /api/v1/notifications/unread-count\n")
    p.add_run("• Event Trigger: ").bold = True
    p.add_run("When an employer updates candidate application status, the backend dispatches a structured in-app notification to the student with notification_type = 'application_status_update'.\n")

    # -------------------------------------------------------------
    # SECTION 3: EMPLOYER PORTAL WORK & DATA FLOW
    # -------------------------------------------------------------
    h1 = doc.add_heading("3. Employer Portal: End-to-End Work & Data Flow", level=1)
    h1.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "The Employer Portal empowers corporate recruiters and hiring managers to publish vacancies, automatically ingest skill requirements, and review candidate applicants pre-sorted by AI match affinity."
    )

    # Step 3.1
    h2 = doc.add_heading("Phase 3.1: Employer Registration & Company Profile", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("RegisterScreen.tsx (Role: 'employer')\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Employer signs up with Company Name, Work Email, and Password. Backend atomically creates User record and linked EmployerProfile record.\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("POST /api/v1/auth/register (role: 'employer')\n")

    # Step 3.2
    h2 = doc.add_heading("Phase 3.2: Job Posting Creation & Vector Indexing", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("EmployerPostingsScreen.tsx (Modal: 'Create New Job Posting')\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Recruiter taps '+ Create Job Posting' -> Fills Title, Job Type ('internship', 'ojt', 'full_time', 'part_time'), Location, Remote/On-site toggle, Required Skills (comma-separated), Preferred Skills, and Description -> Taps 'Publish Posting'.\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("POST /api/v1/jobs/\n")
    p.add_run("• Outgoing Payload:\n")
    p.add_run('  {\n'
               '    "title": "Backend Engineering Intern",\n'
               '    "description": "Looking for passionate Python/FastAPI interns...",\n'
               '    "job_type": "internship",\n'
               '    "location": "Makati / Remote",\n'
               '    "required_skills": ["Python", "FastAPI", "PostgreSQL", "Git"],\n'
               '    "preferred_skills": ["Docker", "Pytest"]\n'
               '  }\n').italic = True
    p.add_run("• Backend Action: ").bold = True
    p.add_run("Skills normalized via ontology; 384-dimensional dense vector embedding generated from synthesized job text representation; posting stored with status = 'active'.\n")

    # Step 3.3
    h2 = doc.add_heading("Phase 3.3: Active Postings Dashboard", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("EmployerPostingsScreen.tsx\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Recruiter views list of all active company postings with applicant count badges (e.g. '12 Applicants'). Recruiter taps 'View Ranked Applicants' on a card to navigate into the candidate review pipeline.\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("GET /api/v1/jobs/my-postings\n")

    # Step 3.4
    h2 = doc.add_heading("Phase 3.4: AI-Ranked Candidate Applicant Review", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("EmployerApplicantsScreen.tsx\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Recruiter views all applicants for the selected posting. Candidates are automatically pre-sorted descending by hybrid AI Match Score (e.g., #1 Juan Dela Cruz - 94% Match, #2 Maria Santos - 82% Match). Each card displays matched skill badges, missing skill alerts, applicant notes, current recruitment status badge, and 'Download Resume' button.\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("GET /api/v1/applications/posting/{posting_id}/applicants\n")
    p.add_run("• Ranked Applicant Object Schema:\n")
    p.add_run('  {\n'
               '    "application_id": "app_55",\n'
               '    "candidate_name": "Juan Dela Cruz",\n'
               '    "candidate_email": "juan@univ.edu",\n'
               '    "status": "under_review",\n'
               '    "match_score": 0.94,\n'
               '    "skill_score": 0.90,\n'
               '    "resume_url": "/api/v1/resumes/files/user_01/resume_uuid.pdf",\n'
               '    "explanation": {\n'
               '      "matched_skills": ["Python", "FastAPI", "PostgreSQL", "Git"],\n'
               '      "missing_skills": ["Docker"],\n'
               '      "summary": "Candidate matches 4 of 5 required skills with strong semantic vector alignment."\n'
               '    }\n'
               '  }\n').italic = True

    # Step 3.5
    h2 = doc.add_heading("Phase 3.5: Recruitment Pipeline Status Progression & Notifications", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)

    p = doc.add_paragraph()
    p.add_run("• Screen / Component: ").bold = True
    p.add_run("EmployerApplicantsScreen.tsx (Action Buttons on candidate card)\n")
    p.add_run("• User Workflow: ").bold = True
    p.add_run("Recruiter taps 1-click status progression buttons: 'Shortlist', 'Schedule Interview', 'Accept Candidate', or 'Reject'.\n")
    p.add_run("• API Endpoint: ").bold = True
    p.add_run("PATCH /api/v1/applications/{application_id}/status\n")
    p.add_run("• Outgoing Payload: ").bold = True
    p.add_run('{ "status": "interview_scheduled" }\n').italic = True
    p.add_run("• Automated Backend Trigger: ").bold = True
    p.add_run("Backend updates Application record AND immediately dispatches an in-app notification to the student's account ('Interview Scheduled: Tech Corp has scheduled an interview for Backend Engineering Intern.').\n")

    # -------------------------------------------------------------
    # SECTION 4: COMPLETE DATA FLOW & API MATRIX TABLE
    # -------------------------------------------------------------
    h1 = doc.add_heading("4. Master API & Data Flow Matrix", level=1)
    h1.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph("Summary of all API routes, HTTP methods, authorization roles, and primary database models:")

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Module"
    hdr[1].text = "HTTP & Route"
    hdr[2].text = "Auth / Role"
    hdr[3].text = "Primary Input / Payload"
    hdr[4].text = "Output & Side Effects"

    api_rows = [
        ("Auth", "POST /auth/register", "Public", "email, password, full_name, role", "Creates User (+ EmployerProfile if employer), returns JWT"),
        ("Auth", "POST /auth/login", "Public", "username (email), password", "Returns JWT access_token + user data"),
        ("Auth", "GET /auth/me/export", "Student / Employer", "None (Bearer JWT)", "RA 10173 complete JSON personal data export"),
        ("Auth", "DELETE /auth/me", "Student / Employer", "None (Bearer JWT)", "RA 10173 permanent cascading delete + disk purge"),
        ("Resumes", "POST /resumes/upload", "Student", "Multipart PDF/DOCX file", "Parses raw text, LLM JSON extract, vector embedding"),
        ("Resumes", "PUT /resumes/{id}/parsed-data", "Student", "parsed_data JSON", "Human verification: re-normalizes skills & vector"),
        ("Resumes", "DELETE /resumes/{id}", "Student", "None", "Deletes resume record and physical file on disk"),
        ("Jobs", "POST /jobs/", "Employer", "title, job_type, skills, desc", "Creates internal posting, builds 384-dim job embedding"),
        ("Jobs", "GET /jobs/my-postings", "Employer", "None", "Returns employer's active postings with applicant counts"),
        ("Matches", "GET /matches/recommendations", "Student", "None", "Hybrid scored feed of internal + external jobs"),
        ("Applications", "POST /applications/", "Student", "job_posting_id, notes", "Creates application in 'submitted' status"),
        ("Applications", "GET /applications/my-applications", "Student", "None", "Returns student's applications with 5-step status"),
        ("Applications", "GET /applications/posting/{id}/applicants", "Employer", "None", "Returns AI-ranked applicant pool descending"),
        ("Applications", "PATCH /applications/{id}/status", "Employer", "status enum", "Transitions status & dispatches student notification"),
        ("Notifications", "GET /notifications/", "Authenticated", "None", "Returns user's in-app notification activity feed"),
        ("Evaluation", "POST /evaluation/run-benchmark", "Coordinator/Admin", "None", "Runs Precision@K, NDCG@K, MRR benchmark"),
    ]

    for mod, route, auth, inp, out in api_rows:
        row = table.add_row().cells
        row[0].text = mod
        row[1].text = route
        row[2].text = auth
        row[3].text = inp
        row[4].text = out

    style_table(table, header_bg="1E293B", col_widths=[0.9, 1.8, 1.1, 1.6, 2.1])
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 5: STATE MACHINES & ENUMS
    # -------------------------------------------------------------
    h1 = doc.add_heading("5. State Machines & Status Enums", level=1)
    h1.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    h2 = doc.add_heading("5.1 Application Status Lifecycle", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)
    doc.add_paragraph(
        "The recruitment lifecycle follows a strict sequential state progression with terminal states:"
    )
    doc.add_paragraph(
        "• submitted: Initial state upon student application.\n"
        "• under_review: Recruiter has opened and viewed the candidate profile.\n"
        "• shortlisted: Candidate marked as high-potential match for interview pool.\n"
        "• interview_scheduled: Candidate invited for technical/behavioral interview.\n"
        "• accepted (Terminal Success): Offer extended and accepted.\n"
        "• rejected (Terminal): Application declined.\n"
        "• withdrawn (Terminal): Student voluntarily withdrew application."
    )

    h2 = doc.add_heading("5.2 Resume Parsing Status Lifecycle", level=2)
    h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)
    doc.add_paragraph(
        "• uploaded  →  parsing  →  parsed (Awaiting student verification)  →  active (Verified & matched)  →  archived"
    )

    # -------------------------------------------------------------
    # SECTION 6: UI/UX IMPROVEMENT BLUEPRINT & ACTIONABLE CHECKLIST
    # -------------------------------------------------------------
    h1 = doc.add_heading("6. UI/UX Improvement Blueprint & Redesign Checklist", level=1)
    h1.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "To elevate the mobile experience from functional prototype to an industry-grade, visually stunning flagship product, "
        "the following concrete UI/UX improvements are recommended for implementation across the mobile app codebase:"
    )

    checklist_items = [
        ("🎨 Design Tokens & Theme System", [
            "• Replace hardcoded hex colors with a centralized Theme / Token system (src/theme/colors.ts).",
            "• Implement a cohesive Dark Glassmorphism palette: Background (#090D16), Surface (#111827), Elevated Surface (#1F2937), Primary Indigo (#6366F1), Emerald Success (#10B981), Amber Warning (#F59E0B), Rose Danger (#EF4444).",
            "• Typography hierarchy: Load Google Fonts (e.g. Outfit or Inter) with dedicated display styles for headings, badges, and metadata labels."
        ]),
        ("📱 Student Screen Improvements", [
            "• ResumeUploadScreen: Add an interactive animated upload drop zone with file type icons, real-time upload progress bar, and pulsating AI parsing skeleton loader.",
            "• ResumeReviewScreen: Add categorized skill pills (Languages, Frameworks, Cloud, Databases) with a quick-tag search bar for seamless skill additions.",
            "• RecommendationFeedScreen: Implement interactive swipeable job cards with a circular gradient Match Score Gauge (e.g., 94%), match breakdown accordion, and sticky Filter/Search bar.",
            "• ApplicationTrackerScreen: Replace text status with an animated vertical/horizontal recruitment timeline stepper with pulsing active nodes and status timestamps."
        ]),
        ("🏢 Employer Screen Improvements", [
            "• EmployerPostingsScreen: Redesign job cards with mini-applicant avatars, quick-toggle active/pause switches, and a streamlined multi-step 'Post a Job' modal with skill tag autocompletion.",
            "• EmployerApplicantsScreen: Add sorting/filtering toolbar (Sort by: Highest AI Match, Most Recent, Status) and an expandable candidate dossier view displaying full resume text and skill match comparison side-by-side."
        ]),
        ("⚡ Micro-Interactions, Feedback & Empty States", [
            "• Skeleton Loaders: Replace generic ActivityIndicator spinners with shimmer skeleton placeholders matching card layouts.",
            "• Empty States: Design custom SVG/vector illustrations with contextual action buttons for 'No Recommendations Yet', 'No Applications Yet', and 'No Applicants Yet'.",
            "• Haptic Feedback: Add subtle Expo Haptics (Haptics.selectionAsync(), Haptics.notificationAsync()) on button taps, status updates, and successful uploads."
        ]),
        ("🛡️ Offline Resilience & Error Handling", [
            "• Toast Notification Banner: Add a floating toast banner for network connection alerts, API error messages, and action confirmations.",
            "• Pull-to-Refresh: Implement RefreshControl across all feed and list screens with smooth indicator animations."
        ])
    ]

    for cat_title, bullets in checklist_items:
        h2 = doc.add_heading(cat_title, level=2)
        h2.runs[0].font.color.rgb = RGBColor(67, 56, 202)
        for b in bullets:
            p = doc.add_paragraph(b)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)

    # Output file
    output_path = Path(r"c:\Users\ajcj1\OneDrive\Desktop\hunting\Ai-hunt\SkillMatch_AI_Work_and_Data_Flow_Specification.docx")
    doc.save(str(output_path))
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_document()
