import io
import os
from typing import Tuple, Optional
import puremagic
import pdfplumber
import docx
from app.core.config import settings


class FileValidationError(Exception):
    """Raised when file fails security or format validation."""
    pass


class ExtractionError(Exception):
    """Raised when text cannot be extracted from a valid document (e.g. scanned image)."""
    pass


def validate_file_magic_bytes(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    Validates file integrity and true MIME type using magic header bytes (puremagic),
    protecting against disguised executables or spoofed extensions.

    Returns:
        Tuple of (verified_mime_type, file_extension)
    """
    if len(file_bytes) == 0:
        raise FileValidationError("Uploaded file is empty (0 bytes).")

    if len(file_bytes) > settings.MAX_UPLOAD_SIZE_BYTES:
        raise FileValidationError(
            f"File size exceeds the maximum limit of {settings.MAX_UPLOAD_SIZE_BYTES / (1024 * 1024):.1f}MB."
        )

    _, ext = os.path.splitext(filename.lower())
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise FileValidationError(
            f"Unsupported file extension '{ext}'. Allowed extensions: {settings.ALLOWED_EXTENSIONS}"
        )

    try:
        matches = puremagic.magic_string(file_bytes)
        if not matches:
            raise FileValidationError("Unable to verify file type signatures.")
        
        detected_mime = matches[0].mime_type
    except Exception as e:
        # Fallback inspection for docx zip container headers if puremagic returns generic zip
        if ext == ".docx" and file_bytes[:4] == b"PK\x03\x04":
            detected_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext == ".pdf" and file_bytes[:4] == b"%PDF":
            detected_mime = "application/pdf"
        else:
            raise FileValidationError(f"File signature verification failed: {str(e)}")

    # Check against allowed MIME types
    if detected_mime not in settings.ALLOWED_MIME_TYPES:
        # Check standard pdf/docx magic byte fallbacks
        if ext == ".pdf" and file_bytes[:4] == b"%PDF":
            detected_mime = "application/pdf"
        elif ext == ".docx" and file_bytes[:4] == b"PK\x03\x04":
            detected_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            raise FileValidationError(
                f"Invalid file MIME type '{detected_mime}'. The file does not match allowed PDF or DOCX formats."
            )

    return detected_mime, ext


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes using pdfplumber.
    Raises ExtractionError if document is scanned or contains no extractable text.
    """
    extracted_text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if len(pdf.pages) == 0:
                raise ExtractionError("PDF contains no pages.")
            
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    extracted_text_parts.append(page_text.strip())

        full_text = "\n\n".join(extracted_text_parts).strip()
        
        if not full_text:
            raise ExtractionError(
                "No readable text found in PDF. Scanned images and photo resumes are not supported. "
                "Please upload a text-based PDF or DOCX document."
            )
        
        return full_text
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError(f"Error extracting text from PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        full_text = "\n".join(paragraphs).strip()
        if not full_text:
            raise ExtractionError("Document contains no readable text.")

        return full_text
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError(f"Error extracting text from DOCX: {str(e)}")


def extract_document_text(file_bytes: bytes, filename: str) -> Tuple[str, str, str]:
    """
    Validates file magic bytes and extracts clean text.

    Returns:
        Tuple of (extracted_text, verified_mime_type, file_extension)
    """
    mime_type, ext = validate_file_magic_bytes(file_bytes, filename)
    
    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        text = extract_text_from_docx(file_bytes)
    else:
        raise FileValidationError(f"Unsupported format: {ext}")

    return text, mime_type, ext
